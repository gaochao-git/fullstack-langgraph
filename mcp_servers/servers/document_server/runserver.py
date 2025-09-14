#!/usr/bin/env python3
"""
Document Generation MCP Server
文档生成MCP服务器
支持生成Word、PDF等格式的文档
"""

import json
import logging
import os
import uuid
from datetime import datetime
from typing import Dict, Any, Optional, List, Tuple
from fastmcp import FastMCP
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from ..common.base_config import MCPServerConfig

# 配置日志
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# 创建MCP服务器实例
mcp = FastMCP("Document Generation Server")

# 加载配置
config = MCPServerConfig('document_server')

# 获取文档存储目录
DOCUMENT_DIR = config.get('document_dir', '/tmp/mcp_documents')
os.makedirs(DOCUMENT_DIR, exist_ok=True)

def setup_document_styles(doc: Document):
    """设置文档样式"""
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # 设置标题样式
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = '黑体'
        heading_style.font.size = Pt(16 - i * 2)
        heading_style.font.bold = True
        heading_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)
    
    # 设置列表样式
    list_style = doc.styles['List Bullet']
    list_style.font.name = '宋体'
    list_style.font.size = Pt(12)

def parse_markdown_table(lines: List[str], start_idx: int) -> Tuple[Any, int]:
    """解析Markdown表格，返回表格对象和结束索引"""
    # 寻找表格的开始和结束
    table_lines = []
    i = start_idx
    
    while i < len(lines):
        line = lines[i].strip()
        if '|' in line:
            table_lines.append(line)
            i += 1
        else:
            break
    
    if len(table_lines) < 2:  # 至少需要表头和分隔线
        return None, start_idx
    
    # 解析表格
    rows = []
    for line in table_lines:
        # 跳过分隔线
        if all(c in '-|' for c in line.replace(' ', '')):
            continue
        # 分割单元格
        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        rows.append(cells)
    
    return rows, i - 1

def add_table_to_doc(doc: Document, table_data: List[List[str]]):
    """将表格数据添加到文档"""
    if not table_data:
        return
    
    # 创建表格
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    
    # 设置表格样式
    try:
        table.style = 'Light Grid Accent 1'  # 更美观的表格样式
    except:
        table.style = 'Table Grid'  # 备用样式
    
    # 填充数据
    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            # 清空单元格
            cell.text = ''
            # 添加段落
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(cell_text)
            run.font.name = '宋体'
            run.font.size = Pt(12)
            
            # 设置对齐方式
            if j == 0:  # 第一列左对齐
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            else:  # 其他列居中
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 第一行（表头）特殊处理
            if i == 0:
                run.bold = True
                run.font.size = Pt(13)
                # 设置表头背景色（如果支持）
                try:
                    from docx.oxml import parse_xml
                    from docx.oxml.ns import nsdecls
                    shading_elm = parse_xml(r'<w:shd {} w:fill="E0E0E0"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                except:
                    pass
    
    # 自动调整列宽
    table.autofit = True

def parse_markdown_to_docx(doc: Document, content: str):
    """将Markdown内容解析并添加到Word文档"""
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        stripped_line = line.strip()
        
        if not stripped_line:
            doc.add_paragraph()  # 空行
            i += 1
            continue
        
        # 检查是否是表格
        if '|' in stripped_line and i + 1 < len(lines) and '|' in lines[i + 1]:
            table_data, end_idx = parse_markdown_table(lines, i)
            if table_data:
                add_table_to_doc(doc, table_data)
                i = end_idx + 1
                continue
        
        # 一级标题
        if stripped_line.startswith('# '):
            heading = doc.add_heading(stripped_line[2:], level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # 二级标题
        elif stripped_line.startswith('## '):
            doc.add_heading(stripped_line[3:], level=2)
        # 三级标题
        elif stripped_line.startswith('### '):
            doc.add_heading(stripped_line[4:], level=3)
        # 无序列表
        elif stripped_line.startswith('- ') or stripped_line.startswith('* '):
            text = stripped_line[2:]
            para = doc.add_paragraph(text, style='List Bullet')
            para.paragraph_format.left_indent = Inches(0.5)
        # 有序列表
        elif len(stripped_line) > 2 and stripped_line[0].isdigit() and stripped_line[1:3] == '. ':
            text = stripped_line[3:]
            para = doc.add_paragraph(text, style='List Number')
            para.paragraph_format.left_indent = Inches(0.5)
        # 普通段落
        else:
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Pt(24)  # 首行缩进两个字符
            # 处理内联格式
            parts = stripped_line.split('**')
            for i_part, part in enumerate(parts):
                if i_part % 2 == 1:  # 粗体
                    run = para.add_run(part)
                    run.bold = True
                    run.font.name = '宋体'
                else:
                    # 处理斜体
                    subparts = part.split('*')
                    for j, subpart in enumerate(subparts):
                        run = para.add_run(subpart)
                        if j % 2 == 1:  # 斜体
                            run.italic = True
                        run.font.name = '宋体'
        
        i += 1

@mcp.tool()
async def generate_word_document(
    title: str,
    content: str,
    template_type: str = "general",
    author: str = "AI Assistant"
) -> str:
    """生成Word文档
    
    Args:
        title: 文档标题
        content: 文档内容（支持Markdown格式）
        template_type: 模板类型 (general/report/summary/proposal)
        author: 文档作者
    
    Returns:
        JSON格式的文档信息，包含文件ID、路径等
    """
    try:
        # 创建文档
        doc = Document()
        
        # 设置文档样式
        setup_document_styles(doc)
        
        # 设置文档属性
        doc.core_properties.author = author
        doc.core_properties.title = title
        doc.core_properties.created = datetime.now()
        
        # 根据模板类型设置样式
        if template_type == "report":
            # 正式报告样式（公文格式）
            # 标题
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.font.name = '黑体'
            title_run.font.size = Pt(18)
            title_run.bold = True
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.space_after = Pt(24)
            
            # 横线
            doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 日期
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(datetime.now().strftime("%Y年%m月%d日"))
            date_run.font.name = '宋体'
            date_run.font.size = Pt(14)
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_para.paragraph_format.space_after = Pt(24)
            
        elif template_type == "summary":
            # 总结样式（年终总结格式）
            # 标题
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.font.name = '黑体'
            title_run.font.size = Pt(22)
            title_run.bold = True
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.space_before = Pt(36)
            title_para.paragraph_format.space_after = Pt(36)
            
            # 日期（右对齐）
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(datetime.now().strftime("%Y年%m月%d日"))
            date_run.font.name = '宋体'
            date_run.font.size = Pt(12)
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_para.paragraph_format.space_after = Pt(24)
            
        else:
            # 通用样式
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.font.name = '黑体'
            title_run.font.size = Pt(16)
            title_run.bold = True
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.space_after = Pt(18)
        
        # 解析并添加内容
        parse_markdown_to_docx(doc, content)
        
        # 生成文件信息
        file_id = str(uuid.uuid4())
        filename = f"{file_id}.docx"
        filepath = os.path.join(DOCUMENT_DIR, filename)
        
        # 保存文档
        doc.save(filepath)
        
        # 返回文档信息
        result = {
            "status": "success",
            "file_id": file_id,
            "title": title,
            "filename": f"{title}.docx",
            "filepath": filepath,
            "size": os.path.getsize(filepath),
            "created_at": datetime.now().isoformat(),
            "template_type": template_type,
            "download_url": f"/documents/{file_id}/download"
        }
        
        logger.info(f"文档生成成功: {filepath}")
        return json.dumps(result, ensure_ascii=False, indent=2)
        
    except Exception as e:
        logger.error(f"生成文档失败: {e}")
        return json.dumps({
            "status": "error",
            "error": str(e)
        }, ensure_ascii=False)

@mcp.tool()
async def generate_annual_summary(
    year: int,
    employee_name: str,
    department: str,
    achievements: List[str],
    challenges: List[str],
    next_year_plans: List[str]
) -> str:
    """生成年终总结文档
    
    Args:
        year: 年份
        employee_name: 员工姓名
        department: 部门
        achievements: 主要成就列表
        challenges: 面临的挑战列表
        next_year_plans: 下一年计划列表
    
    Returns:
        JSON格式的文档信息
    """
    # 构建年终总结内容（公文格式）
    content = f"""**基本信息**

**姓名**：{employee_name}
**部门**：{department}

## 一、工作概述

{year}年度，本人在{department}工作，认真履行岗位职责，积极完成各项工作任务。现将本年度工作情况总结如下：

## 二、主要成就

本年度取得的主要工作成绩包括：

"""
    
    for i, achievement in enumerate(achievements, 1):
        content += f"{i}. {achievement}\n"
    
    content += f"""
## 三、存在的问题与改进措施

工作中遇到的主要困难和挑战：

"""
    
    for i, challenge in enumerate(challenges, 1):
        content += f"{i}. {challenge}\n"
    
    content += f"""
## 四、工作体会

通过本年度的工作实践，本人深刻体会到：团队协作的重要性、持续学习的必要性以及主动沟通在工作中的关键作用。在今后的工作中，将进一步提升专业能力，更好地服务于部门和公司的发展。

## 五、{year+1}年工作计划

"""
    
    for i, plan in enumerate(next_year_plans, 1):
        content += f"{i}. {plan}\n"
    
    content += f"""
## 六、结语

感谢各级领导和同事们的支持与帮助。在新的一年里，本人将以更加饱满的工作热情，为{department}的发展做出更大贡献。

"""
    
    # 生成文档
    title = f"{employee_name}_{year}年度工作总结"
    return await generate_word_document(title, content, "summary", employee_name)

@mcp.tool()
async def generate_official_document(
    title: str,
    document_number: str,
    recipient: str,
    subject: str,
    content: str,
    issuer: str = "办公室",
    cc_list: List[str] = None
) -> str:
    """生成公文格式文档
    
    Args:
        title: 文档标题（如：关于XXX的通知）
        document_number: 文号（如：〔2024〕1号）
        recipient: 主送单位
        subject: 主题词
        content: 正文内容（支持Markdown）
        issuer: 发文单位
        cc_list: 抄送单位列表
    
    Returns:
        JSON格式的文档信息
    """
    try:
        # 创建文档
        doc = Document()
        setup_document_styles(doc)
        
        # 文档头部（发文单位）
        header_para = doc.add_paragraph()
        header_run = header_para.add_run(issuer)
        header_run.font.name = '方正小标宋'
        header_run.font.size = Pt(22)
        header_run.font.bold = True
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_para.paragraph_format.space_after = Pt(12)
        
        # 文号
        if document_number:
            number_para = doc.add_paragraph()
            number_run = number_para.add_run(document_number)
            number_run.font.name = '仿宋'
            number_run.font.size = Pt(16)
            number_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            number_para.paragraph_format.space_after = Pt(24)
        
        # 标题
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.name = '方正小标宋'
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Pt(12)
        title_para.paragraph_format.space_after = Pt(24)
        
        # 主送单位
        if recipient:
            recipient_para = doc.add_paragraph()
            recipient_run = recipient_para.add_run(f"{recipient}：")
            recipient_run.font.name = '仿宋'
            recipient_run.font.size = Pt(16)
            recipient_para.paragraph_format.space_after = Pt(12)
        
        # 正文
        parse_markdown_to_docx(doc, content)
        
        # 发文单位（落款）
        doc.add_paragraph()  # 空行
        issuer_para = doc.add_paragraph()
        issuer_run = issuer_para.add_run(issuer)
        issuer_run.font.name = '仿宋'
        issuer_run.font.size = Pt(16)
        issuer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        issuer_para.paragraph_format.right_indent = Pt(48)
        
        # 日期
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(datetime.now().strftime("%Y年%m月%d日"))
        date_run.font.name = '仿宋'
        date_run.font.size = Pt(16)
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_para.paragraph_format.right_indent = Pt(48)
        
        # 抄送
        if cc_list:
            doc.add_paragraph()  # 空行
            doc.add_paragraph('_' * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
            cc_para = doc.add_paragraph()
            cc_text = f"抄送：{'、'.join(cc_list)}"
            cc_run = cc_para.add_run(cc_text)
            cc_run.font.name = '仿宋'
            cc_run.font.size = Pt(14)
        
        # 主题词
        if subject:
            subject_para = doc.add_paragraph()
            subject_run = subject_para.add_run(f"主题词：{subject}")
            subject_run.font.name = '仿宋'
            subject_run.font.size = Pt(14)
            
        # 生成文件
        file_id = str(uuid.uuid4())
        filename = f"{file_id}.docx"
        filepath = os.path.join(DOCUMENT_DIR, filename)
        
        doc.save(filepath)
        
        result = {
            "status": "success",
            "file_id": file_id,
            "title": title,
            "filename": f"{title}.docx",
            "filepath": filepath,
            "size": os.path.getsize(filepath),
            "created_at": datetime.now().isoformat(),
            "template_type": "official",
            "download_url": f"/documents/{file_id}/download"
        }
        
        logger.info(f"公文生成成功: {filepath}")
        return json.dumps(result, ensure_ascii=False, indent=2)
        
    except Exception as e:
        logger.error(f"生成公文失败: {e}")
        return json.dumps({
            "status": "error",
            "error": str(e)
        }, ensure_ascii=False)

@mcp.tool()
async def list_generated_documents(limit: int = 10) -> str:
    """列出最近生成的文档
    
    Args:
        limit: 返回文档数量限制
    
    Returns:
        JSON格式的文档列表
    """
    try:
        files = []
        for filename in os.listdir(DOCUMENT_DIR):
            if filename.endswith('.docx'):
                filepath = os.path.join(DOCUMENT_DIR, filename)
                stat = os.stat(filepath)
                files.append({
                    'filename': filename,
                    'size': stat.st_size,
                    'created': datetime.fromtimestamp(stat.st_ctime).isoformat()
                })
        
        # 按创建时间排序
        files.sort(key=lambda x: x['created'], reverse=True)
        files = files[:limit]
        
        return json.dumps({
            "status": "success",
            "documents": files,
            "total": len(files)
        }, ensure_ascii=False, indent=2)
        
    except Exception as e:
        logger.error(f"列出文档失败: {e}")
        return json.dumps({
            "status": "error",
            "error": str(e)
        }, ensure_ascii=False)

@mcp.tool()
async def generate_table_document(
    title: str,
    description: str,
    table_data: Dict[str, List[Any]],
    summary: str = ""
) -> str:
    """生成包含表格的文档
    
    Args:
        title: 文档标题
        description: 表格描述
        table_data: 表格数据，格式为 {"headers": ["列1", "列2"], "rows": [["数据1", "数据2"], ...]}
        summary: 总结说明
    
    Returns:
        JSON格式的文档信息
    """
    try:
        # 构建Markdown内容
        content = f"{description}\n\n"
        
        # 构建表格
        if table_data and "headers" in table_data and "rows" in table_data:
            headers = table_data["headers"]
            rows = table_data["rows"]
            
            # 表头
            content += "| " + " | ".join(headers) + " |\n"
            # 分隔线
            content += "|" + "|".join([" --- " for _ in headers]) + "|\n"
            # 数据行
            for row in rows:
                content += "| " + " | ".join(str(cell) for cell in row) + " |\n"
        
        if summary:
            content += f"\n{summary}"
        
        # 使用report模板生成文档
        return await generate_word_document(title, content, "report", "系统")
        
    except Exception as e:
        logger.error(f"生成表格文档失败: {e}")
        return json.dumps({
            "status": "error",
            "error": str(e)
        }, ensure_ascii=False)

@mcp.tool()
async def get_document_content(file_id: str) -> str:
    """获取文档内容预览（返回文档基本信息）
    
    Args:
        file_id: 文档ID
    
    Returns:
        JSON格式的文档信息
    """
    try:
        filepath = os.path.join(DOCUMENT_DIR, f"{file_id}.docx")
        if not os.path.exists(filepath):
            return json.dumps({
                "status": "error",
                "error": "文档不存在"
            }, ensure_ascii=False)
        
        doc = Document(filepath)
        
        # 提取文档信息
        paragraphs = []
        for para in doc.paragraphs[:10]:  # 只预览前10段
            if para.text.strip():
                paragraphs.append(para.text.strip())
        
        return json.dumps({
            "status": "success",
            "file_id": file_id,
            "filepath": filepath,
            "preview": paragraphs,
            "total_paragraphs": len(doc.paragraphs),
            "properties": {
                "author": doc.core_properties.author,
                "title": doc.core_properties.title,
                "created": doc.core_properties.created.isoformat() if doc.core_properties.created else None
            }
        }, ensure_ascii=False, indent=2)
        
    except Exception as e:
        logger.error(f"获取文档内容失败: {e}")
        return json.dumps({
            "status": "error",
            "error": str(e)
        }, ensure_ascii=False)

# 运行服务器
if __name__ == "__main__":
    # 获取端口
    port = config.get('port', 3006)
    logger.info(f"Document MCP Server starting on port {port}")
    logger.info(f"Documents stored in: {DOCUMENT_DIR}")
    logger.info(f"Access at: http://localhost:{port}/sse/")
    
    # 使用 HTTP 传输方式运行
    mcp.run(transport="streamable-http", host="0.0.0.0", port=port)