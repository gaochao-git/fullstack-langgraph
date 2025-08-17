#!/usr/bin/env python3
"""
Document Generation MCP Server
文档生成MCP服务器
支持生成Word、PDF等格式的文档
"""

import json
import logging
import os
import uuid
from datetime import datetime
from typing import Dict, Any, Optional, List
from fastmcp import FastMCP
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from base_config import MCPServerConfig

# 配置日志
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# 创建MCP服务器实例
mcp = FastMCP("Document Generation Server")

# 加载配置
config = MCPServerConfig('document_server')

# 获取文档存储目录
DOCUMENT_DIR = config.get('document_dir', '/tmp/mcp_documents')
os.makedirs(DOCUMENT_DIR, exist_ok=True)

def parse_markdown_to_docx(doc: Document, content: str):
    """将Markdown内容解析并添加到Word文档"""
    lines = content.split('\n')
    current_list = None
    
    for line in lines:
        line = line.strip()
        
        if not line:
            continue
            
        # 一级标题
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        # 二级标题
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        # 三级标题
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        # 无序列表
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            doc.add_paragraph(text, style='List Bullet')
        # 有序列表
        elif line[0].isdigit() and line[1:3] == '. ':
            text = line[3:]
            doc.add_paragraph(text, style='List Number')
        # 普通段落
        else:
            para = doc.add_paragraph()
            # 处理内联格式
            parts = line.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 1:  # 粗体
                    para.add_run(part).bold = True
                else:
                    # 处理斜体
                    subparts = part.split('*')
                    for j, subpart in enumerate(subparts):
                        if j % 2 == 1:  # 斜体
                            para.add_run(subpart).italic = True
                        else:
                            para.add_run(subpart)

@mcp.tool()
async def generate_word_document(
    title: str,
    content: str,
    template_type: str = "general",
    author: str = "AI Assistant"
) -> str:
    """生成Word文档
    
    Args:
        title: 文档标题
        content: 文档内容（支持Markdown格式）
        template_type: 模板类型 (general/report/summary/proposal)
        author: 文档作者
    
    Returns:
        JSON格式的文档信息，包含文件ID、路径等
    """
    try:
        # 创建文档
        doc = Document()
        
        # 设置文档属性
        doc.core_properties.author = author
        doc.core_properties.title = title
        doc.core_properties.created = datetime.now()
        
        # 根据模板类型设置样式
        if template_type == "report":
            # 正式报告样式
            title_para = doc.add_heading(title, 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加日期
            date_para = doc.add_paragraph(datetime.now().strftime("%Y年%m月%d日"))
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # 空行
            
        elif template_type == "summary":
            # 总结样式
            title_para = doc.add_heading(title, 0)
            
            # 添加作者和日期信息
            info_para = doc.add_paragraph()
            info_para.add_run(f"作者：{author} | ").italic = True
            info_para.add_run(f"日期：{datetime.now().strftime('%Y-%m-%d')}").italic = True
            
            doc.add_paragraph()  # 空行
            
        else:
            # 通用样式
            doc.add_heading(title, 0)
            doc.add_paragraph()
        
        # 解析并添加内容
        parse_markdown_to_docx(doc, content)
        
        # 生成文件信息
        file_id = str(uuid.uuid4())
        filename = f"{file_id}.docx"
        filepath = os.path.join(DOCUMENT_DIR, filename)
        
        # 保存文档
        doc.save(filepath)
        
        # 返回文档信息
        result = {
            "status": "success",
            "file_id": file_id,
            "title": title,
            "filename": f"{title}.docx",
            "filepath": filepath,
            "size": os.path.getsize(filepath),
            "created_at": datetime.now().isoformat(),
            "template_type": template_type,
            "download_url": f"/documents/{file_id}/download"
        }
        
        logger.info(f"文档生成成功: {filepath}")
        return json.dumps(result, ensure_ascii=False, indent=2)
        
    except Exception as e:
        logger.error(f"生成文档失败: {e}")
        return json.dumps({
            "status": "error",
            "error": str(e)
        }, ensure_ascii=False)

@mcp.tool()
async def generate_annual_summary(
    year: int,
    employee_name: str,
    department: str,
    achievements: List[str],
    challenges: List[str],
    next_year_plans: List[str]
) -> str:
    """生成年终总结文档
    
    Args:
        year: 年份
        employee_name: 员工姓名
        department: 部门
        achievements: 主要成就列表
        challenges: 面临的挑战列表
        next_year_plans: 下一年计划列表
    
    Returns:
        JSON格式的文档信息
    """
    # 构建年终总结内容
    content = f"""# {year}年度工作总结

## 基本信息
- **姓名**：{employee_name}
- **部门**：{department}
- **日期**：{datetime.now().strftime('%Y年%m月%d日')}

## 一、工作概述

本年度在{department}部门工作期间，我认真履行岗位职责，积极配合团队完成各项工作任务。现将{year}年度的工作情况总结如下。

## 二、主要成就

"""
    
    for i, achievement in enumerate(achievements, 1):
        content += f"{i}. {achievement}\n"
    
    content += """
## 三、面临的挑战与解决方案

在工作过程中，我也遇到了一些挑战：

"""
    
    for i, challenge in enumerate(challenges, 1):
        content += f"{i}. {challenge}\n"
    
    content += """
## 四、经验与收获

通过这一年的工作，我深刻认识到：
- 团队协作的重要性
- 持续学习和技能提升的必要性
- 主动沟通和问题解决能力的价值

## 五、下一年工作计划

展望{next_year}年，我将重点做好以下工作：

""".format(next_year=year+1)
    
    for i, plan in enumerate(next_year_plans, 1):
        content += f"{i}. {plan}\n"
    
    content += f"""
## 六、结语

感谢领导和同事们在过去一年中给予的支持和帮助。在新的一年里，我将继续努力，为{department}和公司的发展贡献自己的力量。

---
*{employee_name}*  
*{datetime.now().strftime('%Y年%m月%d日')}*
"""
    
    # 生成文档
    title = f"{employee_name}_{year}年度工作总结"
    return await generate_word_document(title, content, "summary", employee_name)

@mcp.tool()
async def list_generated_documents(limit: int = 10) -> str:
    """列出最近生成的文档
    
    Args:
        limit: 返回文档数量限制
    
    Returns:
        JSON格式的文档列表
    """
    try:
        files = []
        for filename in os.listdir(DOCUMENT_DIR):
            if filename.endswith('.docx'):
                filepath = os.path.join(DOCUMENT_DIR, filename)
                stat = os.stat(filepath)
                files.append({
                    'filename': filename,
                    'size': stat.st_size,
                    'created': datetime.fromtimestamp(stat.st_ctime).isoformat()
                })
        
        # 按创建时间排序
        files.sort(key=lambda x: x['created'], reverse=True)
        files = files[:limit]
        
        return json.dumps({
            "status": "success",
            "documents": files,
            "total": len(files)
        }, ensure_ascii=False, indent=2)
        
    except Exception as e:
        logger.error(f"列出文档失败: {e}")
        return json.dumps({
            "status": "error",
            "error": str(e)
        }, ensure_ascii=False)

@mcp.tool()
async def get_document_content(file_id: str) -> str:
    """获取文档内容预览（返回文档基本信息）
    
    Args:
        file_id: 文档ID
    
    Returns:
        JSON格式的文档信息
    """
    try:
        filepath = os.path.join(DOCUMENT_DIR, f"{file_id}.docx")
        if not os.path.exists(filepath):
            return json.dumps({
                "status": "error",
                "error": "文档不存在"
            }, ensure_ascii=False)
        
        doc = Document(filepath)
        
        # 提取文档信息
        paragraphs = []
        for para in doc.paragraphs[:10]:  # 只预览前10段
            if para.text.strip():
                paragraphs.append(para.text.strip())
        
        return json.dumps({
            "status": "success",
            "file_id": file_id,
            "filepath": filepath,
            "preview": paragraphs,
            "total_paragraphs": len(doc.paragraphs),
            "properties": {
                "author": doc.core_properties.author,
                "title": doc.core_properties.title,
                "created": doc.core_properties.created.isoformat() if doc.core_properties.created else None
            }
        }, ensure_ascii=False, indent=2)
        
    except Exception as e:
        logger.error(f"获取文档内容失败: {e}")
        return json.dumps({
            "status": "error",
            "error": str(e)
        }, ensure_ascii=False)

# 运行服务器
if __name__ == "__main__":
    # 从配置获取端口
    port = config._server_info.get('port', 3006)
    logger.info(f"Document MCP Server starting on port {port}")
    logger.info(f"Documents stored in: {DOCUMENT_DIR}")
    logger.info(f"Access at: http://localhost:{port}/sse/")
    
    # 使用 HTTP 传输方式运行
    mcp.run(transport="streamable-http", host="0.0.0.0", port=port)